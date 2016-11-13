import sys
import os
from PyQt4 import QtCore, QtGui,uic
from docx import Document
from docx.shared import Inches
from faker import Factory
from Crypto.Hash import SHA256
import zipfile
import tarfile


if hasattr(sys, '_MEIPASS'):
    qtCreatorFile = os.path.join(sys._MEIPASS, "gui.ui")
else:
    qtCreatorFile= "gui.ui"

Ui_MainWindow, QtBaseClass = uic.loadUiType(qtCreatorFile)    # Load the UI
    
class MyApp(QtGui.QMainWindow, Ui_MainWindow):
    def __init__(self):
        QtGui.QMainWindow.__init__(self)
        Ui_MainWindow.__init__(self)
        self.setupUi(self)
        self.Enter.clicked.connect(self.Event) #Once the Enter button on the GUI is pressed
        # this code class the Event function. 
         
    
    # Event grabs the information provided by the drop down boxes in the gui. 
    def Event(self):
        exten = self.exten.currentText()
        compress = self.compress.currentText()
        encrypt = self.encrypt.currentText()
        self.extension(exten,encrypt,compress)
        self.compression(compress,exten)
        window.close()
    
    # Compression function: Calls all the compression type functions    
    def compression(self,compress,exten):
        if(compress == "ZIP"):
            self.comp_zip(exten)
        elif(compress =="TAR"):
            self.comp_tar(exten)
        else:
            print "Some modules are still being prepared"
    
    # Recieves a ssn or ccn and hashes it using SHA256 and returns that value.     
    def encryption(self,val):
        encrypt = SHA256.new()
        encrypt.update(val)        
        return encrypt.hexdigest()
    
    #Generates the files by calling individual file creation functions. 
    def extension(self,exten,encrypt,compress):
        #easy_file = ["xlsx","csv","xls","txt","conf"]
        #if(exten in easy_file):
         #   self.generate(exten,compress,encrypt)
        if(exten == "doc" or exten == "docx"):
            self.doc(exten,compress,encrypt)
    #Generates random SSNs
    def ssn(self):
        fake = Factory.create()
        return fake.ssn()
    #Generates random CCNs using visa as its numbering scheme
    def ccn(self):
        fake = Factory.create()
        return fake.credit_card_number(card_type="visa")
    
    #A genearlized file generation function.
    def generate(self,exten,compress,encrypt):
        f = open(('002x134uz.'+str(exten)),'w')
        if(encrypt == "None"):
            for i in range(100):
                f.write(self.ssn()+'\n')
                f.write(self.ccn()+'\n')
           
        else:
            for i in range(100):
                f.write(self.encryption(self.ssn())+'\n')
                f.write(self.encryption(self.ccn())+'\n')
        f.closed  
        
        if(compress != "None"):
            self.compression(compress,exten)
        
    #Creates doc and docx files using the Documents Library.   
    def doc(self,exten,compress,encrypt):
        document = Document()

        if(encrypt == "None"):
            for i in range(0,100):
                document.add_paragraph(self.ssn())
                document.add_paragraph(self.ccn())
        else:
            for i in range(0,100):
                document.add_paragraph(self.ssn())
                document.add_paragraph(self.ccn())
                
        document.save('002x134uz.'+str(exten))
        
        if(compress != "None"):
            self.compression(exten,compress)
        
    
    # Compression function for zips    
    def comp_zip(self,exten):
        doc = '002x134uz.' + str(exten) 
        zf = zipfile.ZipFile('002x134uz.zip',mode='w')
        try:
            zf.write(doc)
        finally:
            zf.close()
        self.cleanup(exten)
    
    # Compression function for tar files    
    def comp_tar(self,exten):
        doc = '002x134uz.' + str(exten) 
        out = tarfile.open('002x134uz.tar',mode='w')
        try:
            out.add(doc)
        finally:
            out.close()
        self.cleanup(exten)
    
    #This removes a file.         
    def cleanup(self,exten):
        try:
            os.remove('002x134uz.'+str(exten))
        except:
            return
             
        
if __name__ == "__main__":
    app = QtGui.QApplication(sys.argv)
    window = MyApp()
    window.show()
    sys.exit(app.exec_())
    
