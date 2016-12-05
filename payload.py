from rstr import xeger
from docx import Document
from Crypto.Hash import SHA256
import zipfile
import tarfile
import os
import shutil 
import xlsxwriter

# Encrypted takes a string an create 
# a hex dump using SHA256 hashing.
def encrypted(val):
        encrypt = SHA256.new()
        encrypt.update(val)        
        return encrypt.hexdigest()
                    
# reg_generate takes a regular expression
# and generates a string based off the regular expression
def reg_generate(text):
    string = xeger(text)
    return string  

# extension takes text and 
# calls all the different file generators.
def extensions(text):
    file_doc(text)
    file_txt(text)
    #file_pdf(text)
    file_xlsx(text)
    file_csv(text)
# compression takes a file and compresses it to all
# the compression files.    
def compression(file_):
    comp_zip(file_)
    comp_tar(file_)
    comp_bz2(file_)
    #comp_rar(file_)
    comp_tar_gz(file_)
    
    #This moves all the text based files to the 
    # FILES folder with in the directory. 
    shutil.move(file_,"./FILES")

# comp_zip takes a file name and writes it as a 
# zip file using the same name with just a .zip file
# extension.     
def comp_zip(file_):
    doc = file_
    zip_f = file_ + '.zip'
    zf = zipfile.ZipFile( zip_f,mode='w')
    try:
        zf.write(doc)
    finally:
        zf.close()
    # This moves the created zip file to the subdirectory. 
    shutil.move(zip_f,"./FILES/Compressed_File")

# comp_tar takes a file name and writes it as a 
# tar file using the same name with just a .tar file
# extension.             
def comp_tar(file_):
    doc =  file_
    tar_f = file_ + '.tar'
    out = tarfile.open( tar_f,mode='w')
    try:
        out.add(doc)
    finally:
        out.close()
     # This moves the created zip file to the subdirectory. 
    shutil.move(tar_f,"./FILES/Compressed_File")   
    
# comp_bz2 takes a file name and writes it as a 
# tarball with a bz2 compression write. 
def comp_bz2(file_):
    doc = file_
    bz2 = file_ + ".tar.bz2"
    bz2_f = tarfile.open( bz2, mode="w:bz2")
    try:
        bz2_f.add(doc)
    finally:
        bz2_f.close()

    shutil.move(bz2,"./FILES/Compressed_File")  

# comp_tar_gz takes a file name and compresses the file
# to a tarball gzip file. 
def comp_tar_gz(file_):
    doc = file_
    gz = file_ + ".tar.gz"
    tar_gz = tarfile.open(gz, mode='w:gz')
    try:
        tar_gz.add(doc)
    finally:
        tar_gz.close()  
        
    shutil.move(gz,"./FILES/Compressed_File")
    

# file_doc creates both doc and docx files. 
# This takes regular expression list.   
def file_doc(text):
    # This creates the 4 different documents. 
    docx = Document()
    doc = Document()
    docxen = Document()
    docen = Document()    
    #-------------------------
    #----This runs the regex list and generates 
    #----data based off the regex. 
    for i in text:
        for k in range(100):
            doc.add_paragraph(reg_generate(i))
            docx.add_paragraph(reg_generate(i))
            docen.add_paragraph(encrypted(reg_generate(i)))
            docxen.add_paragraph(encrypted(reg_generate(i)))            
    
    #----This flushes the files out to the disk
    docx.save("unencrypted.docx")
    doc.save("unencrypted.doc")
    docxen.save("encrypted.docx")
    docen.save("encrypted.doc") 
    #------------------------------

    #----These compress the newly created files. 
    compression("unencrypted.docx")
    compression("unencrypted.doc")
    compression("encrypted.docx")
    compression("encrypted.doc") 
    #------------------------------------

# Generates txt files writing string returned from the regular expression.
# The data is either encrypted or not. Then the files are compressed. 
def file_txt(text):
    unencrypted = open("unencrypted.txt",'w')
    encrypted_ = open("encrypted.txt",'w')
    
    for i in text:
            for k in range(100):
                unencrypted.write(str(reg_generate(i)) + '\n')
                encrypted_.write(str(encrypted(reg_generate(i)))+ '\n')
    
    unencrypted.close()
    encrypted_.close()
    
    compression("unencrypted.txt")
    compression("encrypted.txt")   
# Generates csv files writing string returned from the regular expression.
# The data is either encrypted or not. Then the files are compressed. 
def file_csv(text):
    unencrypted = open("unencrypted.csv",'w')
    encrypted_ = open("encrypted.csv",'w')    
    for i in text:
        for k in range(100):
            unencrypted.write(str(reg_generate(i)) + '\n')
            encrypted_.write(str(encrypted(reg_generate(i)))+ '\n')
    unencrypted.close()
    encrypted_.close()
    compression("unencrypted.csv")
    compression("encrypted.csv")

def file_xlsx(text):
    #First to create an xlsx or xls file you create the overall file.
    #Second: Then you add a sheet to that workbook.  
    
    workbook_xu = xlsxwriter.Workbook('unencrypted.xlsx')
    worksheet_xu = workbook_xu.add_worksheet()
    
    workbook_xe = xlsxwriter.Workbook('encrypted.xlsx')
    worksheet_xe = workbook_xe.add_worksheet()   
    
    workbook_u = xlsxwriter.Workbook('unencrypted.xls')
    worksheet_u = workbook_u.add_worksheet()
   
    workbook_e = xlsxwriter.Workbook('encrypted.xls')
    worksheet_e = workbook_e.add_worksheet()        
    
    #The you can write to each worksheet the generated data. 
    for i in text:
        for k in range(100):
            worksheet_xu.write(k,0,str(reg_generate(i)))
            worksheet_xe.write(k,0,str(encrypted(reg_generate(i))))
            worksheet_u.write(k,0,str(reg_generate(i)))
            worksheet_e.write(k,0,str(encrypted(reg_generate(i))))
    
    #Flush the files to the disk
    workbook_xu.close()
    workbook_xe.close()
   
    workbook_u.close()
    workbook_e.close()
   
    # Compress the files    
    compression("unencrypted.xlsx")
    compression("unencrypted.xls")
    compression("encrypted.xls")
    compression("encrypted.xlsx")     

def main():
   
    print "To use hard coded data just press enter."
    reg = str(raw_input("Enter a string or a regular expression: "))
   
    inp =[]
   
    path = "./FILES"
    path2 = "/Compressed_File"
   
    if not os.path.exists(path):
        os.makedirs(path)
        os.makedirs(path + path2)
    
    else:
        shutil.rmtree(path)
        os.makedirs(path)
        os.makedirs(path + path2)

    #
    while reg != "":
        inp.append(reg)
        print "Hit Enter on an empty line to stop."
        reg = str(raw_input("Enter a string or a regular expression: "))

    if len(inp) == 0:
        #Expression for social security numbers
        inp.append("[0-9][0-9][0-9]\-[0-9][0-9]\-[0-9][0-9][0-9][0-9]") 
   
        #Expression for visa credit card numbers
        inp.append("^4[0-9]{12}(?:[0-9]{3})?$")
        
        #Expression for master card number
        inp.append("^5[1-5][0-9]{14}$")
        
        #A list of regular expressions can be found in regular_expression.txt
    
    
    extensions(inp)   
main()